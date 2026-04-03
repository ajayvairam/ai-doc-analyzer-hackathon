"""
AI-Powered Document Analysis & Extraction — Processor Module
Handles text extraction (python-docx / OCR.space) and Gemini AI analysis.
"""

import os
import re
import json
import base64
import io
import requests
from dotenv import load_dotenv

load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
OCR_SPACE_API_KEY = os.getenv("OCR_SPACE_API_KEY")

# ---------------------------------------------------------------------------
# Text Extraction
# ---------------------------------------------------------------------------

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                paragraphs.append(" | ".join(row_text))

    return "\n".join(paragraphs)


def extract_text_via_ocr(file_base64: str, file_type: str) -> str:
    """Extract text from PDF or image using OCR.space Free API, with retry & engine fallback."""
    if not OCR_SPACE_API_KEY:
        raise ValueError("OCR_SPACE_API_KEY is not set in environment variables.")

    import time
    
    engines_to_try = [1, 2]  # Engine 1 is usually more stable for free tier
    last_error = None

    for engine in engines_to_try:
        for attempt in range(2):
            payload = {
                "apikey": OCR_SPACE_API_KEY,
                "base64Image": file_base64,  # Now includes the full data URL from browser
                "language": "eng",
                "isOverlayRequired": False,
                "detectOrientation": True,
                "scale": True,
                "OCREngine": engine,
            }

            try:
                response = requests.post(
                    "https://api.ocr.space/parse/image",
                    data=payload,
                    timeout=30,
                )
                response.raise_for_status()
                result = response.json()

                if result.get("IsErroredOnProcessing"):
                    error_msg = result.get("ErrorMessage", ["Unknown OCR error"])
                    # If it's a known error, we can still fall back to the next engine
                    raise RuntimeError(f"OCR.space error: {error_msg}")

                # Concatenate parsed text from all pages
                parsed_results = result.get("ParsedResults", [])
                texts = [pr.get("ParsedText", "") for pr in parsed_results]
                return "\n".join(texts).strip()
                
            except requests.exceptions.RequestException as e:
                last_error = f"OCR Request failed (Engine {engine}): {e}"
                time.sleep(2)  # Wait before retry
            except Exception as e:
                last_error = f"OCR Processing failed (Engine {engine}): {e}"
                time.sleep(2)  # Wait before retry
                
    raise RuntimeError(f"OCR extraction failed after retries. Last error: {last_error}")


def extract_text(file_base64: str, file_type: str) -> str:
    """
    Route to the correct extraction method based on file type.
    - 'docx' → python-docx
    - 'pdf' / 'image' → OCR.space
    """
    if file_type == "docx":
        # Extract the actual base64 string from the data URL
        raw_b64 = file_base64.split(",")[1] if "," in file_base64 else file_base64
        file_bytes = base64.b64decode(raw_b64)
        return extract_text_from_docx(file_bytes)
    elif file_type in ("pdf", "image"):
        return extract_text_via_ocr(file_base64, file_type)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


# ---------------------------------------------------------------------------
# Gemini AI Analysis
# ---------------------------------------------------------------------------

SYSTEM_INSTRUCTION = """You are a professional document analyst AI. 
Analyze the provided document text thoroughly.

You MUST return ONLY valid JSON with NO additional text, no markdown fences, no explanation. 

The JSON structure must be EXACTLY:
{
  "status": "success",
  "summary": "A concise 2-3 sentence summary of the document's content and purpose.",
  "entities": {
    "names": ["List of person names found in the document"],
    "dates": ["List of dates found in the document"],
    "organizations": ["List of company/organization names found"],
    "amounts": ["List of monetary amounts/currency values found, e.g. '$5,000', '€1,200', 'INR 50,000'"]
  },
  "sentiment": "Positive or Neutral or Negative"
}

Rules:
- Extract ALL monetary amounts and currency values into the "amounts" array. This is critical.
- If no entities of a category are found, use an empty array [].
- The summary must be exactly 2-3 sentences.
- Sentiment must be exactly one of: "Positive", "Neutral", or "Negative".
- Do NOT wrap the JSON in markdown code fences.
- Do NOT include any text outside the JSON object.
"""


def analyze_with_gemini(text: str, file_name: str) -> dict:
    """
    Send extracted text to Gemini AI for structured analysis.
    Tries multiple models with retry logic if quota is exceeded.
    Returns a dict with status, summary, entities, and sentiment.
    """
    if not GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY is not set in environment variables.")

    import time
    from google import genai
    from google.genai import types

    client = genai.Client(api_key=GEMINI_API_KEY)

    # Models to try in order — each has its own separate quota
    MODEL_CANDIDATES = [
        "gemini-2.5-flash",
        "gemini-2.0-flash",
        "gemini-2.0-flash-lite",
        "gemini-2.5-flash-lite",
    ]

    full_prompt = f"""{SYSTEM_INSTRUCTION}

Analyze the following document.

Document Name: {file_name}

--- DOCUMENT TEXT START ---
{text[:15000]}
--- DOCUMENT TEXT END ---

Return ONLY the JSON analysis as specified above."""

    last_error = None
    raw = None

    for model_name in MODEL_CANDIDATES:
        for attempt in range(2):  # 2 attempts per model
            try:
                response = client.models.generate_content(
                    model=model_name,
                    contents=full_prompt,
                    config=types.GenerateContentConfig(
                        temperature=0.1,
                        max_output_tokens=1024,
                    ),
                )
                raw = response.text.strip()
                break  # Success — exit retry loop
            except Exception as e:
                last_error = e
                error_str = str(e)
                if "429" in error_str or "quota" in error_str.lower() or "RESOURCE_EXHAUSTED" in error_str:
                    if attempt == 0:
                        time.sleep(6)  # Brief wait before retry
                        continue
                    else:
                        break  # Move to next model
                else:
                    raise  # Non-quota error — raise immediately
        else:
            continue  # All retries for this model exhausted, try next
        break  # Success — exit model loop
    else:
        raise RuntimeError(f"All Gemini models quota exceeded. Last error: {last_error}")

    # Strip markdown code fences if model wraps them
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    raw = raw.strip()

    try:
        result = json.loads(raw)
    except json.JSONDecodeError:
        # Attempt to find JSON object in the response
        match = re.search(r"\{.*\}", raw, re.DOTALL)
        if match:
            result = json.loads(match.group())
        else:
            raise RuntimeError(f"Gemini returned invalid JSON: {raw[:500]}")

    # Ensure required keys exist with defaults
    result.setdefault("status", "success")
    result.setdefault("summary", "No summary available.")
    result.setdefault("entities", {})
    result["entities"].setdefault("names", [])
    result["entities"].setdefault("dates", [])
    result["entities"].setdefault("organizations", [])
    result["entities"].setdefault("amounts", [])
    result.setdefault("sentiment", "Neutral")

    return result
